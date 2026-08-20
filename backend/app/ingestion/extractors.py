"""File-type extractors. Each takes a local Path and returns an ExtractedDocument.

Status vocabulary matches the ingestion report requirement in the plan:
  ok / partial / unsupported / failed  (dedup + "current vs stale" happen later in the
  pipeline, not here).
"""

from __future__ import annotations

import re
from pathlib import Path

import fitz  # pymupdf
import olefile
import pdfplumber
from docx import Document as DocxDocument

from app.ingestion.extracted import ExtractedDocument, ExtractedPage, ExtractedTable

MIN_CHARS_PER_PAGE_FOR_TEXT_LAYER = 20
OCR_RENDER_DPI = 300


def _configure_tesseract() -> bool:
    """Points pytesseract at the configured binary. Returns False (and leaves OCR
    disabled) if no path is configured or pytesseract isn't importable, so callers
    can degrade gracefully instead of crashing ingestion."""
    from app.config import get_settings

    settings = get_settings()
    if not settings.tesseract_cmd:
        return False
    try:
        import pytesseract

        pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd
        return True
    except ImportError:
        return False


def _ocr_image_bytes(png_bytes: bytes) -> str:
    import io

    import pytesseract
    from PIL import Image

    with Image.open(io.BytesIO(png_bytes)) as img:
        return pytesseract.image_to_string(img)


def extract_pdf(path: Path, ocr_available: bool = False) -> ExtractedDocument:
    warnings: list[str] = []
    pages: list[ExtractedPage] = []
    scanned_pages: list[int] = []
    ocr_pages: list[int] = []
    tesseract_ready = ocr_available and _configure_tesseract()

    try:
        fitz_doc = fitz.open(path)
    except Exception as e:
        return ExtractedDocument(status="failed", reason=f"Could not open PDF: {e}")

    # pdfplumber gives more reliable table extraction than pymupdf's raw text.
    try:
        plumber_doc = pdfplumber.open(path)
    except Exception as e:
        plumber_doc = None
        warnings.append(f"Table extraction unavailable (pdfplumber failed to open: {e})")

    for i, fpage in enumerate(fitz_doc):
        page_number = i + 1
        text = fpage.get_text()
        char_count = len(text.strip())

        if char_count < MIN_CHARS_PER_PAGE_FOR_TEXT_LAYER:
            scanned_pages.append(page_number)
            if not tesseract_ready:
                # No text layer and no OCR engine configured: page contributes nothing.
                pages.append(ExtractedPage(page_number=page_number, text=""))
                continue
            try:
                pix = fpage.get_pixmap(dpi=OCR_RENDER_DPI)
                text = _ocr_image_bytes(pix.tobytes("png"))
                ocr_pages.append(page_number)
            except Exception as e:
                warnings.append(f"OCR failed on page {page_number}: {e}")
                text = ""

        headings = _detect_pdf_headings(fpage)
        tables: list[ExtractedTable] = []
        if plumber_doc is not None:
            try:
                ppage = plumber_doc.pages[i]
                for raw_table in ppage.extract_tables():
                    cleaned = [[(c or "").strip() for c in row] for row in raw_table]
                    if cleaned:
                        tables.append(ExtractedTable(page_number=page_number, rows=cleaned))
            except Exception as e:
                warnings.append(f"Table extraction failed on page {page_number}: {e}")

        pages.append(
            ExtractedPage(page_number=page_number, text=text, headings=headings, tables=tables)
        )

    fitz_doc.close()
    if plumber_doc is not None:
        plumber_doc.close()

    total_chars = sum(p.char_count for p in pages)
    unrecovered_pages = [p for p in scanned_pages if p not in ocr_pages]

    if total_chars == 0:
        reason = "No extractable text layer on any page (scanned/image-only PDF)"
        if tesseract_ready:
            reason += "; OCR was attempted but produced no usable text."
        else:
            reason += ". Install Tesseract and set TESSERACT_CMD to index this file's content."
        return ExtractedDocument(status="unsupported", reason=reason, pages=pages, warnings=warnings)

    if ocr_pages:
        warnings.append(
            f"{len(ocr_pages)} page(s) had no text layer and were recovered via OCR "
            f"(pages: {ocr_pages[:20]}{'...' if len(ocr_pages) > 20 else ''}). OCR text can contain "
            "recognition errors; treat citations to these pages with extra care."
        )
    if unrecovered_pages:
        warnings.append(
            f"{len(unrecovered_pages)} page(s) had no text layer and could not be recovered "
            f"(pages: {unrecovered_pages[:20]}{'...' if len(unrecovered_pages) > 20 else ''})."
        )

    if scanned_pages:
        # Any page that relied on OCR (or is still missing) makes the whole
        # document 'partial' rather than 'ok' — an honest confidence signal,
        # since OCR text is less reliable than a native PDF text layer.
        return ExtractedDocument(status="partial", reason="; ".join(warnings), pages=pages, warnings=warnings)

    return ExtractedDocument(status="ok", reason=None, pages=pages, warnings=warnings)


def _detect_pdf_headings(fpage) -> list[tuple[str, int]]:
    """Heuristic heading detection from font size: lines whose font size is
    meaningfully larger than the page's modal (body-text) font size are treated
    as headings, ranked by size (0 = largest)."""
    try:
        raw = fpage.get_text("dict")
    except Exception:
        return []

    sizes: list[float] = []
    lines_with_size: list[tuple[str, float]] = []
    for block in raw.get("blocks", []):
        for line in block.get("lines", []):
            spans = line.get("spans", [])
            if not spans:
                continue
            text = "".join(s.get("text", "") for s in spans).strip()
            if not text:
                continue
            size = max(s.get("size", 0) for s in spans)
            sizes.append(size)
            lines_with_size.append((text, size))

    if not sizes:
        return []

    sizes_sorted = sorted(sizes)
    body_size = sizes_sorted[len(sizes_sorted) // 2]  # median ~= body text size
    distinct_larger = sorted({s for s in sizes if s > body_size + 0.5}, reverse=True)
    rank_by_size = {s: i for i, s in enumerate(distinct_larger)}

    headings = []
    for text, size in lines_with_size:
        if size in rank_by_size and len(text) < 150:
            headings.append((text, rank_by_size[size]))
    return headings


def extract_docx(path: Path) -> ExtractedDocument:
    try:
        doc = DocxDocument(str(path))
    except Exception as e:
        return ExtractedDocument(status="failed", reason=f"Could not open DOCX: {e}")

    text_parts: list[str] = []
    headings: list[tuple[str, int]] = []
    tables: list[ExtractedTable] = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        text_parts.append(t)
        style = (para.style.name or "") if para.style else ""
        m = re.match(r"Heading (\d)", style)
        if m:
            headings.append((t, int(m.group(1)) - 1))

    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if rows:
            tables.append(ExtractedTable(page_number=None, rows=rows))

    full_text = "\n".join(text_parts)
    if not full_text.strip() and not tables:
        return ExtractedDocument(status="failed", reason="DOCX contained no extractable text or tables.")

    page = ExtractedPage(page_number=1, text=full_text, headings=headings, tables=tables)
    return ExtractedDocument(status="ok", reason=None, pages=[page])


_OLE_UNICODE_RUN = re.compile(rb"(?:[\x20-\x7E\xA0-\xFF]\x00){4,}")
_OLE_NOISE_PATTERNS = re.compile(
    r"^(Normal|Heading \d|Times New Roman|Arial|Symbol|Default Paragraph Font|"
    r"Table Normal|No List)$",
    re.IGNORECASE,
)


def extract_legacy_doc(path: Path) -> ExtractedDocument:
    """Best-effort extraction for pre-2007 binary .doc files without LibreOffice/
    antiword: scan every OLE stream for runs of UTF-16LE-encoded printable text.
    This recovers most body text but loses structure (headings, tables, page
    breaks) and may include some style-name noise. Always reported as 'partial'."""
    try:
        if not olefile.isOleFile(str(path)):
            return ExtractedDocument(
                status="failed", reason="File is not a valid OLE compound document (corrupt or not a real .doc)."
            )
        ole = olefile.OleFileIO(str(path))
    except Exception as e:
        return ExtractedDocument(status="failed", reason=f"Could not open as OLE file: {e}")

    candidate_streams = [s for s in ole.listdir() if "WordDocument" in s or "Table" in "".join(s)]
    if not candidate_streams:
        candidate_streams = ole.listdir()

    seen_lines: set[str] = set()
    ordered_lines: list[str] = []
    for stream_path in candidate_streams:
        try:
            data = ole.openstream(stream_path).read()
        except Exception:
            continue
        for match in _OLE_UNICODE_RUN.finditer(data):
            try:
                decoded = match.group(0).decode("utf-16-le", errors="ignore")
            except Exception:
                continue
            for line in re.split(r"[\r\v\x00]+", decoded):
                line = line.strip()
                if len(line) < 3:
                    continue
                if _OLE_NOISE_PATTERNS.match(line):
                    continue
                if line not in seen_lines:
                    seen_lines.add(line)
                    ordered_lines.append(line)
    ole.close()

    full_text = "\n".join(ordered_lines)
    if len(full_text) < 200:
        return ExtractedDocument(
            status="failed",
            reason=(
                "Legacy .doc heuristic byte-scan recovered too little text to be usable "
                f"({len(full_text)} chars). This file needs a proper .doc parser (e.g. "
                "LibreOffice headless --convert-to docx) to index reliably."
            ),
        )

    page = ExtractedPage(page_number=1, text=full_text)
    return ExtractedDocument(
        status="partial",
        reason=(
            "Extracted via heuristic byte-scan of the legacy .doc binary format (no "
            "LibreOffice/antiword available on this machine). Headings, tables, and page "
            "structure are lost, word order within recovered runs may be imperfect, and some "
            "boilerplate style-name noise may remain. Re-ingest with LibreOffice installed "
            "(or after converting the source to .docx/.pdf) for full fidelity."
        ),
        pages=[page],
        warnings=["Legacy .doc: structure-free heuristic extraction."],
    )


def extract_image(path: Path, ocr_available: bool = False) -> ExtractedDocument:
    if not (ocr_available and _configure_tesseract()):
        return ExtractedDocument(
            status="unsupported",
            reason=(
                "Image file requires OCR to extract any text and Tesseract is not installed "
                "on this machine. Set TESSERACT_CMD in .env once installed to index this file."
            ),
        )
    try:
        text = _ocr_image_bytes(path.read_bytes())
    except Exception as e:
        return ExtractedDocument(status="failed", reason=f"OCR failed: {e}")

    if len(text.strip()) < 10:
        return ExtractedDocument(
            status="partial",
            reason=(
                "OCR ran but recovered very little text; the image may be mostly a diagram "
                "with little running text (e.g. a plumbing schematic). Recovered text, if any, "
                "was still indexed."
            ),
            pages=[ExtractedPage(page_number=1, text=text)] if text.strip() else [],
        )

    return ExtractedDocument(
        status="partial",  # OCR output always carries lower confidence than native text.
        reason="Text recovered via OCR; may contain recognition errors.",
        pages=[ExtractedPage(page_number=1, text=text)],
    )


def extract_indd(path: Path) -> ExtractedDocument:
    return ExtractedDocument(
        status="unsupported",
        reason=(
            "Adobe InDesign (.indd) has no reliable open-source parser. Export this file to "
            "PDF or DOCX from InDesign and re-upload to include its content in the knowledge base."
        ),
    )


EXTENSION_MAP = {
    ".pdf": "pdf",
    ".doc": "doc",
    ".docx": "docx",
    ".jpg": "image",
    ".jpeg": "image",
    ".png": "image",
    ".indd": "indd",
}

_MAGIC_SIGNATURES: list[tuple[bytes, str]] = [
    (b"%PDF-", "pdf"),
    (b"PK\x03\x04", "zip_ooxml"),          # .docx/.xlsx/.pptx are zip containers
    (b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1", "ole"),  # legacy .doc/.xls/.ppt
    (b"\xff\xd8\xff", "image"),
    (b"\x89PNG\r\n\x1a\n", "image"),
]


def sniff_file_type(path: Path) -> str | None:
    """Detect actual file type from magic bytes, independent of extension.
    Several files in the source corpus have extensions that don't match their
    real content (e.g. a PDF saved with a .doc extension) — see the ingestion
    report's 'extension_mismatch' notes."""
    try:
        with open(path, "rb") as f:
            head = f.read(8)
    except Exception:
        return None
    for sig, kind in _MAGIC_SIGNATURES:
        if head.startswith(sig):
            if kind == "zip_ooxml":
                return "docx"  # only .docx is a supported OOXML type here
            return kind
    return None


def classify_extension(path: Path) -> str:
    return EXTENSION_MAP.get(path.suffix.lower(), "unknown")


def resolve_file_type(path: Path) -> tuple[str, str | None]:
    """Returns (effective_type, mismatch_note). Trusts the magic-byte sniff over
    the file extension whenever they disagree, since content parsing depends on
    real format, not the filename."""
    ext_type = classify_extension(path)
    sniffed = sniff_file_type(path)
    if sniffed is None:
        return ext_type, None
    if ext_type == "indd":
        # .indd has no reliable magic-byte signature we check for; trust extension.
        return ext_type, None
    if sniffed != ext_type:
        note = (
            f"File extension ({path.suffix}) suggests '{ext_type}' but file content is "
            f"actually '{sniffed}'; processed as '{sniffed}'."
        )
        return sniffed, note
    return ext_type, None


def extract(path: Path, ocr_available: bool = False) -> tuple[str, ExtractedDocument, str | None]:
    file_type, mismatch_note = resolve_file_type(path)
    if file_type == "pdf":
        doc = extract_pdf(path, ocr_available=ocr_available)
    elif file_type == "docx":
        doc = extract_docx(path)
    elif file_type == "doc":
        doc = extract_legacy_doc(path)
    elif file_type == "image":
        doc = extract_image(path, ocr_available=ocr_available)
    elif file_type == "indd":
        doc = extract_indd(path)
    else:
        doc = ExtractedDocument(status="unsupported", reason=f"Unrecognized file extension: {path.suffix}")

    if mismatch_note:
        doc.warnings.append(mismatch_note)
    return file_type, doc, mismatch_note
